"""
Build a CTO-facing DOCX report from a data.json artefact.

Source-agnostic. Reads `source`, `window_label`, `summary`, `dims` from
the JSON and generates prose dynamically based on the observed numbers.
Sections that depend on a missing dimension (e.g. no `action` field on
the source) are automatically skipped.

Usage:
    python3 build_docx.py --data reports/Prompt_Security_7d/data.json

Output lands next to the input JSON (reports/<slug>_<window>/<slug>_CTO_Report_<window>.docx).
"""
from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# palette (matches charts + pptx)
NAVY = RGBColor(0x0B, 0x1F, 0x3A)
CORAL = RGBColor(0xE8, 0x43, 0x4F)
AMBER = RGBColor(0xF2, 0xA9, 0x50)
TEAL = RGBColor(0x4F, 0xB8, 0x9E)
SLATE = RGBColor(0x88, 0x99, 0xA6)
INK = RGBColor(0x1A, 0x1F, 0x2E)
MUTED = RGBColor(0x55, 0x5F, 0x70)
CREAM_HEX = "F5F1E8"


def _slugify(name: str) -> str:
    s = re.sub(r"[^0-9A-Za-z]+", "_", name.strip())
    return re.sub(r"_+", "_", s).strip("_") or "source"


# ------------------------------------------------------- docx helpers

def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _spacing(p, before=6, after=6, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _run(p, text, bold=False, size=11, color=None, italic=False,
         font="Calibri"):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def _h1(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=18, after=6)
    _run(p, text, bold=True, size=20, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "0B1F3A")
    bottom.set(qn("w:space"), "4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=14, after=4)
    _run(p, text, bold=True, size=14, color=NAVY)


def _h3(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=8, after=2)
    _run(p, text, bold=True, size=12, color=INK)


def _body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    _spacing(p, before=2, after=6, line=1.25)
    _run(p, text, size=11, color=color or INK, italic=italic)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=2, after=2, line=1.25)
    _run(p, "\u25A0 ", size=10, color=NAVY, bold=True)
    _run(p, text, size=11, color=INK)
    return p


def _sub_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.35)
    _spacing(p, before=1, after=1, line=1.25)
    _run(p, "\u2022 ", size=10, color=MUTED, bold=True)
    _run(p, text, size=10.5, color=INK)
    return p


def _page_break(doc):
    p = doc.add_paragraph()
    _spacing(p, before=0, after=0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _stat_card(table, row, col, label, value, color, unit=""):
    cell = table.cell(row, col)
    cell.text = ""
    _set_cell_shading(cell, CREAM_HEX)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, sz, c in (("top", 4, "CCCCCC"), ("bottom", 4, "CCCCCC"),
                        ("right", 4, "CCCCCC")):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), c)
        tcBorders.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    hexc = "{:02X}{:02X}{:02X}".format(*color) if isinstance(color, tuple) \
        else f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" \
        if hasattr(color, "__iter__") else str(color)
    left.set(qn("w:color"), hexc)
    tcBorders.append(left)
    tcPr.append(tcBorders)
    p1 = cell.paragraphs[0]
    _spacing(p1, before=6, after=0)
    _run(p1, str(value), bold=True, size=26, color=NAVY)
    if unit:
        _run(p1, " " + unit, bold=False, size=12, color=MUTED)
    p2 = cell.add_paragraph()
    _spacing(p2, before=0, after=6)
    _run(p2, label, size=10, color=MUTED)


def _add_chart(doc, charts_dir: Path, fn: str, width_in=6.3):
    path = charts_dir / fn
    if not path.exists():
        _body(doc, f"(missing chart: {fn})", color=CORAL)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=6, after=4)
    p.add_run().add_picture(str(path), width=Inches(width_in))


def _caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=8)
    _run(p, text, italic=True, size=9, color=MUTED)


# ------------------------------------------------------- commentary

def _mix_rows(d):
    return (d["queries"].get("per_user_mix_top10", {})
            .get("rows", []) or [])


def _principal_key(mix_rows):
    for r in mix_rows:
        for k in ("user", "src.hostname", "src.ip.address"):
            if k in r:
                return k
    return None


def _principal_label(p_key: Optional[str]) -> str:
    if not p_key:
        return "principals"
    return {"user": "users", "src.hostname": "hosts",
            "src.ip.address": "source IPs"}.get(p_key, p_key)


def _derive_by_principal(mix_rows, p_key):
    acc = defaultdict(int)
    for r in mix_rows:
        acc[r.get(p_key)] += int(r.get("n") or 0)
    return sorted(acc.items(), key=lambda kv: -kv[1])


def _action_top(mix_rows, p_key, action, n=12):
    filt = [r for r in mix_rows
            if str(r.get("action")) == action
            and r.get(p_key) not in (None, "", "null")]
    filt.sort(key=lambda r: -int(r.get("n") or 0))
    return filt[:n]


def _fmt_pct(p: float) -> str:
    if p >= 10:
        return f"{p:.1f}%"
    if p >= 1:
        return f"{p:.2f}%"
    return f"{p:.3f}%"


def _concentration_note(top_share: float) -> str:
    if top_share >= 90:
        return ("Activity is extremely concentrated: a single principal drove "
                "the overwhelming majority of events, which distorts any "
                "baseline calculation until that account is understood.")
    if top_share >= 50:
        return ("Activity is concentrated on a small number of principals. "
                "Tag those principals before deriving a steady-state "
                "baseline.")
    if top_share >= 20:
        return ("Activity is moderately distributed. No single principal "
                "dominates, but the top cohort still accounts for a large "
                "share of volume.")
    return ("Activity is spread evenly across principals. The top cohort "
            "does not dominate, which is the healthier shape for a "
            "steady-state baseline.")


def _intervention_note(pct: float) -> str:
    if pct >= 40:
        return ("The intervention rate is high. Either the policy is "
                "tuned aggressively or the traffic contains a lot of "
                "behaviour the business cares about; both are worth "
                "understanding before promoting or relaxing policies.")
    if pct >= 10:
        return ("The intervention rate is moderate. The policy is doing "
                "real work on a meaningful share of traffic without "
                "blocking the business outright.")
    return ("The intervention rate is low. The policy is in a "
            "predominantly-observe posture, which is the right shape "
            "for a recently-onboarded source or a log-only deployment.")


def _bypass_note(pct: float) -> str:
    if pct < 0.5:
        return ("Bypass events are near-zero. Users are accepting policy "
                "decisions, which is the strongest signal that the policy "
                "is trusted.")
    if pct < 5:
        return ("Bypass events are low but non-trivial. The users doing "
                "the overriding are worth reviewing: low-volume, "
                "high-signal.")
    return ("Bypass events are elevated. A meaningful share of traffic "
            "is opting out of policy decisions; review which policies "
            "and which users are involved.")


# ------------------------------------------------------- sections

def build_header_footer(doc, source, window):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    _spacing(hp, before=0, after=0)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(hp, f"SentinelOne  ", size=9, color=MUTED)
    _run(hp, f"|  {source} {window} activity review",
         size=9, color=NAVY, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    _spacing(fp, before=0, after=0)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, "Prepared by SentinelOne Pre-Sales Engineering  ",
         size=9, color=MUTED)
    _run(fp, "Generated on ", size=9, color=MUTED)
    _run(fp, datetime.now(timezone.utc).strftime("%Y-%m-%d"),
         size=9, color=NAVY, bold=True)


def build_cover(doc, d):
    source = d["source"]
    window = d["window_label"]
    start = d["window_start"][:10]
    end = d["window_end"][:10]
    for _ in range(3):
        p = doc.add_paragraph()
        _spacing(p, before=0, after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=8, after=8)
    _run(p, f"CTO REPORT  |  {source.upper()}",
         bold=True, size=11, color=CORAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=4, after=4)
    _run(p, f"{window} of activity under policy",
         bold=True, size=32, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=24)
    _run(p,
         f"What {source} observed, what we intervened on, "
         "and where to take the insight next.",
         italic=True, size=13, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=4, after=4)
    _run(p, "Window  ", size=11, color=MUTED)
    _run(p, f"{start}  to  {end}  (UTC)",
         bold=True, size=11, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=6)
    _run(p, "Data source  ", size=11, color=MUTED)
    _run(p, f"dataSource.name = '{source}'",
         bold=True, size=11, color=NAVY, font="Consolas")
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=0)
    _run(p,
         "Telemetry pulled live via the SentinelOne Singularity Data "
         "Lake (PowerQuery over the Management Console API).",
         italic=True, size=9, color=MUTED)


def build_stat_row(doc, d, summary, principal_count, window):
    dims = d.get("dims", {})
    p_key = summary.get("top_principal_key")
    t = doc.add_table(rows=1, cols=4)
    t.autofit = False
    widths = [Inches(1.6)] * 4
    for i, w in enumerate(widths):
        t.columns[i].width = w
        for cell in t.columns[i].cells:
            cell.width = w
    _stat_card(t, 0, 0, f"Total events ({window})",
               f"{summary['total']:,}", NAVY)
    if dims.get("action"):
        _stat_card(t, 0, 1, "Blocked events",
                   f"{summary['block']:,}", CORAL)
        _stat_card(t, 0, 2, "Intervention rate",
                   f"{summary['intervention_pct']:.1f}", AMBER, unit="%")
    else:
        # source has no action field; surface the dominant primary
        # key and tenant rank instead of fake zeroes
        by = summary.get("by_action") or {}
        rows = sorted(by.items(), key=lambda kv: -kv[1])
        dom_label = rows[0][0] if rows else "(none)"
        # truncate long event.type labels to fit the card
        if len(dom_label) > 28:
            dom_label = dom_label[:25] + "..."
        _stat_card(t, 0, 1,
                   f"Dominant {summary.get('prim_key', 'value')}",
                   dom_label, SLATE)
        _stat_card(t, 0, 2, "Slices (timeline)",
                   f"{d.get('strategy', {}).get('n_slices', '?')}",
                   AMBER)
    if p_key and principal_count > 0:
        _stat_card(t, 0, 3,
                   f"Distinct {_principal_label(p_key)}",
                   f"{principal_count}", TEAL)
    else:
        rank = summary.get("rank_24h")
        _stat_card(t, 0, 3, "Tenant rank (24h)",
                   f"#{rank}" if rank else "n/a", TEAL)


def build_exec_summary(doc, d):
    _h1(doc, "Executive summary")
    summary = d["summary"]
    source = d["source"]
    window = d["window_label"]
    dims = d.get("dims", {})
    mix = _mix_rows(d)
    p_key = summary.get("top_principal_key") or _principal_key(mix)
    principals = _derive_by_principal(mix, p_key) if p_key else []
    principal_count = len(principals)
    prim_key = summary.get("prim_key")

    # opening sentence adapts to whether we have a principal dim or not
    if p_key and principal_count > 0:
        intro = (f"Over the {window} window ending "
                 f"{d['window_end'][:10]} UTC, "
                 f"dataSource.name='{source}' produced "
                 f"{summary['total']:,} events across "
                 f"{principal_count} distinct "
                 f"{_principal_label(p_key)}. ")
    else:
        intro = (f"Over the {window} window ending "
                 f"{d['window_end'][:10]} UTC, "
                 f"dataSource.name='{source}' produced "
                 f"{summary['total']:,} events. No principal dimension "
                 "(user, host, or IP) was present on this source, so "
                 "this report focuses on volume, timing, and tenant "
                 "context. ")
    # intervention note only makes sense if action is present
    tail = (_intervention_note(summary.get("intervention_pct", 0.0))
            if dims.get("action") else
            "The source does not carry a policy-action field, so this "
            "report treats it as an observability stream rather than "
            "an enforcement surface.")
    _body(doc, intro + tail)

    build_stat_row(doc, d, summary, principal_count, window)

    _h3(doc, "What the CTO should take away")
    if dims.get("action"):
        _bullet(doc,
                f"Block rate of {_fmt_pct(summary['block_pct'])} "
                f"({summary['block']:,} of {summary['total']:,}). "
                "These are events the policy refused outright.")
        if summary.get("modify"):
            _bullet(doc,
                    f"Modify action ran {summary['modify']:,} times "
                    f"({_fmt_pct(100.0 * summary['modify'] / summary['total'])}). "
                    "The system transformed the payload in place instead "
                    "of refusing, letting the user continue.")
        _bullet(doc,
                f"Bypass activity sits at "
                f"{_fmt_pct(summary['bypass_pct'])} "
                f"({summary['bypass']:,} events). "
                + _bypass_note(summary.get("bypass_pct", 0.0)))
    if summary.get("top_user"):
        tu = summary["top_user"]
        _bullet(doc,
                f"{_concentration_note(summary.get('top_share', 0.0))} "
                f"The top principal ({tu.get(p_key)}) drove "
                f"{int(tu['n']):,} events "
                f"({summary['top_share']:.1f}% of volume).")

    # fallbacks when no action / principal: speak to what we do have
    if not dims.get("action") and not summary.get("top_user"):
        by = summary.get("by_action") or {}
        rows = sorted(by.items(), key=lambda kv: -kv[1])
        if rows and prim_key:
            top_label, top_n = rows[0]
            share = 100.0 * top_n / summary["total"] if summary["total"] else 0
            _bullet(doc,
                    f"Dominant {prim_key}: '{top_label}' at "
                    f"{top_n:,} events ({share:.1f}% of the window). "
                    "Use this as the entry point when building "
                    "detections against the source.")
        if summary.get("rank_24h"):
            _bullet(doc,
                    f"In the last 24 hours, {source} ranks "
                    f"#{summary['rank_24h']} among the data sources "
                    "landing in this tenant. The ranking is a useful "
                    "anchor for whether volume is trending or flat.")
        _bullet(doc,
                f"{source} lives in the same data lake as SentinelOne "
                "EDR, identity, network, and cloud telemetry. Any "
                "detection or hunt built on this source can join to "
                "endpoint context in one PowerQuery.")

    _h3(doc, "What we recommend next")
    recs = _recommendations(d)
    for _when, _color, _title, body in recs[:4]:
        _bullet(doc, f"{_title}. {body}")


def build_volume_section(doc, d, charts_dir):
    _h1(doc, "Volume and timing")
    kind = (d["queries"].get("daily_by_action", {})
            .get("slice_kind", "day"))
    kind_adj = {"day": "Daily", "hour": "Hourly",
                "week": "Weekly", "month": "Monthly"}.get(kind,
                                                         kind.title())
    _body(doc,
          f"{kind_adj} event count for {d['source']}, stacked by "
          "policy action where available. The timeline lets you see "
          "whether volume is steady, bursty, or driven by a handful "
          "of incidents during the window.")
    _add_chart(doc, charts_dir, "02_daily_timeline.png", width_in=6.6)
    _caption(doc,
             f"Timeline of {d['source']} events across the window.")

    # derive quick shape analysis
    rows = (d["queries"].get("daily_by_action", {})
            .get("rows", []) or [])
    if rows:
        totals = [int(r.get("matchCount") or 0) for r in rows]
        mx = max(totals) if totals else 0
        mx_idx = totals.index(mx) if totals else 0
        mx_label = rows[mx_idx]["date"] if totals else ""
        tot_sum = sum(totals) or 1
        share = 100 * mx / tot_sum
        _h2(doc, "What the shape tells us")
        if share >= 70:
            _bullet(doc,
                    f"Heavily front- or back-loaded: "
                    f"{mx_label} alone contributed {mx:,} events "
                    f"({share:.1f}% of the window). A single-day spike "
                    "usually points to a demo, a load test, or an "
                    "incident; worth tagging out before deriving a "
                    "steady-state baseline.")
        elif share >= 30:
            _bullet(doc,
                    f"Uneven but not pathological: the busiest slice "
                    f"({mx_label}, {mx:,} events) carried "
                    f"{share:.1f}% of the window. A normal mix of "
                    "bursty and quiet periods.")
        else:
            _bullet(doc,
                    f"Evenly distributed. The busiest slice "
                    f"({mx_label}, {mx:,} events, {share:.1f}% of the "
                    f"window) is only modestly ahead of average; "
                    "this is the steady-state shape.")
        # action stability across days (if action present)
        action_tot = defaultdict(int)
        for r in rows:
            for k, v in (r.get("by_action", {}) or {}).items():
                action_tot[k] += int(v)
        if action_tot:
            primary = max(action_tot.items(), key=lambda kv: kv[1])[0]
            _bullet(doc,
                    f"Dominant action across the window: "
                    f"'{primary}'. Check whether the ratio of '{primary}' "
                    "to the next action stays stable across slices; "
                    "drift is a useful early signal for policy change.")


def build_action_mix(doc, d, charts_dir):
    if not d.get("dims", {}).get("action"):
        return
    _h1(doc, "Policy action mix")
    _body(doc,
          f"Every {d['source']} event terminates in one of a small "
          "set of outcomes. The mix is the most compact answer to "
          "'what is this policy doing for us'.")
    _add_chart(doc, charts_dir, "01_action_mix.png", width_in=6.0)
    total = d["summary"]["total"]
    _caption(doc,
             f"{d['window_label']} distribution of action outcomes across "
             f"{total:,} events.")

    rows = d["queries"]["by_action"]["rows"]
    _h3(doc, "What each action means")
    t = doc.add_table(rows=1, cols=4)
    t.autofit = False
    widths = [Inches(0.9), Inches(1.2), Inches(1.0), Inches(3.3)]
    for i, w in enumerate(widths):
        t.columns[i].width = w
        for c in t.columns[i].cells:
            c.width = w
    hdr = t.rows[0]
    for i, label in enumerate(["Action", "Count", "% of total",
                                "Typical meaning"]):
        cell = hdr.cells[i]
        cell.text = ""
        _set_cell_shading(cell, "0B1F3A")
        _set_cell_border(cell, color="0B1F3A", size=4)
        p = cell.paragraphs[0]
        _spacing(p, before=3, after=3)
        _run(p, label, bold=True, size=10.5,
             color=RGBColor(0xFF, 0xFF, 0xFF))

    meanings = {
        "block": "Event refused by policy; the request was not completed.",
        "modify": "Payload rewritten in place and forwarded (eg. PII redacted).",
        "log": "Event allowed, recorded for audit.",
        "bypass": "A principal explicitly overrode a policy decision.",
        "None": "No action field recorded; typically internal or system traffic.",
        "null": "No action field recorded; typically internal or system traffic.",
    }
    colors = {"block": "FDECEE", "modify": "FEF4E7", "log": "EEF1F3",
              "bypass": "E7F5F1", "None": "F3F4F5", "null": "F3F4F5"}
    rows = sorted(rows, key=lambda r: -int(r.get("n") or 0))
    for r in rows:
        action = str(r.get("action"))
        n = int(r.get("n") or 0)
        row = t.add_row()
        for i, val in enumerate([
            action.upper() if action not in ("None", "null") else "(no action)",
            f"{n:,}",
            f"{100*n/total:.1f}%" if total else "",
            meanings.get(action, "Source-specific outcome."),
        ]):
            cell = row.cells[i]
            cell.width = widths[i]
            cell.text = ""
            _set_cell_shading(cell, colors.get(action, "FFFFFF"))
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            _spacing(p, before=3, after=3)
            _run(p, val, bold=(i == 0), size=10.5, color=INK)


def build_user_section(doc, d, charts_dir):
    summary = d["summary"]
    p_key = summary.get("top_principal_key")
    if not p_key or not summary.get("top_user"):
        return
    _h1(doc, f"Who is driving the traffic")
    _body(doc,
          f"Ranking {_principal_label(p_key)} by volume is how we "
          "identify the accounts that shape the baseline and the "
          "accounts that warrant individual review. The chart uses "
          "a log scale so a dominant principal does not flatten the "
          "long tail.")
    _add_chart(doc, charts_dir, "03_top_users_log.png", width_in=6.6)
    _caption(doc,
             f"Top {_principal_label(p_key)} by {d['window_label']} "
             "event volume (logarithmic x-axis).")

    tu = summary["top_user"]
    _h2(doc, "The concentration story")
    _body(doc,
          f"{tu.get(p_key)} drove {int(tu['n']):,} events "
          f"({summary['top_share']:.1f}% of the window). "
          + _concentration_note(summary.get("top_share", 0.0)))
    if summary.get("top_share", 0.0) >= 50:
        _sub_bullet(doc,
                    "If this was a demo, a test harness, or a scripted "
                    "load: tag it out of baseline calculations before "
                    "moving any policy into enforce.")
        _sub_bullet(doc,
                    "If this was a real principal: pivot into "
                    "SentinelOne EDR/identity for the session history "
                    "and posture of the underlying endpoint.")
        _sub_bullet(doc,
                    "If this cannot be attributed: the anonymous or "
                    "shared-account pattern is itself worth tracking; "
                    "a tag pass here will pay off for every future "
                    "report.")


def build_action_detail_section(doc, d, charts_dir):
    summary = d["summary"]
    dims = d.get("dims", {})
    if not dims.get("action") or not summary.get("top_principal_key"):
        return
    p_key = summary["top_principal_key"]
    _h1(doc, "Risk signals: interventions and overrides")
    _body(doc,
          "Two lenses matter more than the headline numbers: who the "
          "system is intervening on, and who is overriding it. "
          "Intervention tells us where the policy is doing the most "
          "work. Override tells us where principals disagree.")

    _h2(doc, f"Top {_principal_label(p_key)} by BLOCK")
    _add_chart(doc, charts_dir, "04_blocks_by_user.png", width_in=6.6)
    _caption(doc,
             f"{_principal_label(p_key).title()} ranked by block count. "
             "The cohort near the top of the long tail is usually where "
             "human attention pays off.")

    if summary.get("bypass", 0) > 0:
        _h2(doc, f"Top {_principal_label(p_key)} by BYPASS")
        _body(doc,
              f"Bypass is low-volume but highly informative. These are "
              f"the {_principal_label(p_key)} who saw a policy "
              f"decision and chose to override. "
              + _bypass_note(summary.get("bypass_pct", 0.0)))
        _add_chart(doc, charts_dir, "05_bypass_by_user.png", width_in=6.4)
        _caption(doc,
                 f"Top {_principal_label(p_key)} by bypass count. The "
                 "signal to watch over time is concentration.")

    _h2(doc, "Per-principal action mix")
    _body(doc,
          f"Stacked view of the top active {_principal_label(p_key)} "
          "broken out by outcome. This is the clearest way to spot a "
          "mixed-use profile, in which a single principal is being both "
          "blocked and allowed across the window.")
    _add_chart(doc, charts_dir, "06_user_action_mix.png", width_in=6.6)
    _caption(doc,
             f"Action mix per top {_principal_label(p_key)} "
             "(log-scale).")


def build_context_section(doc, d, charts_dir):
    _h1(doc, f"Where {d['source']} fits in the broader picture")
    summary = d["summary"]
    rank = summary.get("rank_24h")
    rank_sentence = (
        f"In the last 24 hours, {d['source']} ranks #{rank} among "
        "the data sources landing in your SentinelOne tenant."
        if rank else
        f"{d['source']} is present in the tenant, though it did not "
        "rank within the top 30 over the last 24 hours.")
    _body(doc,
          f"{rank_sentence} The chart below places {d['source']} in "
          "context alongside EDR, network, identity, and cloud "
          "telemetry so you can see how it contributes to the overall "
          "detection surface.")
    _add_chart(doc, charts_dir, "07_tenant_context.png", width_in=6.6)
    _caption(doc,
             "Data-source volume in the last 24 hours "
             f"({d['source']} highlighted).")

    _h2(doc, "Why this matters to the CTO")
    _bullet(doc,
            f"{d['source']} lives in the same data lake as EDR, "
            "identity, network, and cloud telemetry. Correlation is "
            "native, not a bolt-on.")
    _bullet(doc,
            f"Any detection built on {d['source']} can chain to "
            "SentinelOne endpoint context in a single query; a STAR / "
            "Custom Detection Rule needs no extra plumbing.")
    _bullet(doc,
            "As usage grows, this source's share of daily events will "
            "climb in rank without any ingestion changes.")


def _recommendations(d):
    """Produce data-driven recommendation items. Returns a list of
    (when, color, title, body) tuples."""
    summary = d["summary"]
    dims = d.get("dims", {})
    p_key = summary.get("top_principal_key")
    tu = summary.get("top_user")
    recs = []

    if tu and summary.get("top_share", 0.0) >= 40:
        recs.append((
            "24 hours",
            CORAL,
            f"Resolve the top principal: {tu.get(p_key)}",
            f"A single principal drove "
            f"{summary['top_share']:.1f}% of the window's volume. "
            "Pivot into SentinelOne EDR / identity from this "
            "principal's session to establish whether this is a "
            "scripted harness, a shared account, a misconfigured "
            "integration, or hostile activity. A short investigation "
            "closes the largest single-source risk in this report."))

    if (dims.get("action")
            and summary.get("block_pct", 0.0) >= 10
            and summary.get("bypass_pct", 0.0) < 1.0):
        recs.append((
            "This week",
            AMBER,
            "Promote log-only policies to enforce",
            f"Block rate is "
            f"{_fmt_pct(summary['block_pct'])} and bypass is "
            f"{_fmt_pct(summary['bypass_pct'])}. That combination is "
            "as strong a signal as we ever get that principals accept "
            "policy decisions. Walk each currently-log-only policy "
            "and promote the ones with zero false positives."))

    if dims.get("action") and summary.get("bypass", 0) > 0:
        recs.append((
            "This month",
            TEAL,
            "Correlate bypass with endpoint posture",
            f"Build a Custom Detection Rule that chains a "
            f"{d['source']} bypass event to a SentinelOne alert of "
            "medium-or-higher severity on the same principal within "
            "60 minutes. One rule, two data sources, one unified alert."))

    recs.append((
        "This month",
        NAVY,
        "Enrich ingestion where it is shallow",
        f"Today, {d['source']} events land with a subset of "
        "fields populated. Check whether prompt / payload / content "
        "is null; adding even one content field unlocks content-based "
        "detection and Purple AI hunts against the same data lake."))

    recs.append((
        "Quarterly",
        SLATE,
        "Baseline normal",
        f"Re-run this report on a fixed cadence. The charts are "
        "regenerated from a single JSON artefact, so a 30-day or "
        "90-day view is a parameter flip. Anomalies against the "
        "baseline become the input to the next round of policy "
        "tuning."))

    return recs


def build_recommendations(doc, d):
    _h1(doc, "Recommendations, prioritised")
    for when, color, title, body in _recommendations(d):
        p = doc.add_paragraph()
        _spacing(p, before=10, after=2)
        _run(p, f"[{when}]  ", bold=True, size=11, color=color)
        _run(p, title, bold=True, size=12.5, color=NAVY)
        _body(doc, body)


def build_methodology(doc, d):
    _h1(doc, "Methodology")
    _body(doc,
          "Every number and chart in this report was computed directly "
          "from the SentinelOne Singularity Data Lake using the "
          "Management Console API. No data was moved outside the "
          "tenant. No sampling was applied.")
    _h3(doc, "Data collection")
    _bullet(doc,
            f"Source: dataSource.name = '{d['source']}', excluding the "
            "internal logVolume bookkeeping tag.")
    _bullet(doc,
            f"Window: {d['window_start'][:16]} to "
            f"{d['window_end'][:16]} UTC "
            f"({d['window_label']}).")
    _bullet(doc,
            "Transport: POST /sdl/v2/api/queries (Long Running Query "
            "API) authenticated with Bearer token scoped to the account.")
    strat = d.get("strategy", {})
    _bullet(doc,
            "Aggregations executed in parallel (max 3 concurrent) to "
            "respect the per-user rate limit. "
            f"Timeline rendered via "
            f"{strat.get('n_slices', '?')} "
            f"{strat.get('slice_kind', 'day')}-size "
            "client-side slices (the engine does not expose a native "
            "timebucket function).")
    _h3(doc, "Provenance")
    _body(doc,
          f"A single JSON artefact (reports/{d['slug']}_"
          f"{d['window_label']}/data.json) captures the raw "
          "aggregation output: query text, match counts, elapsed "
          "time, and every returned row. The charts and this document "
          "are pure rendering steps over that artefact, so the "
          "numbers can be reproduced at any time.")
    _h3(doc, "Dimensions observed")
    dims = d.get("dims", {})
    present = [k for k, v in dims.items() if v]
    absent = [k for k, v in dims.items() if not v]
    _bullet(doc,
            f"Present: {', '.join(present) if present else '(none)'}")
    _bullet(doc,
            f"Absent: {', '.join(absent) if absent else '(none)'}")


# ------------------------------------------------------- main

def build_report(data_path: Path) -> Path:
    d = json.loads(data_path.read_text())
    charts_dir = data_path.parent / "charts"
    slug = d.get("slug") or _slugify(d.get("source", "source"))
    window = d.get("window_label", "window")
    out_path = data_path.parent / f"{slug}_CTO_Report_{window}.docx"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    build_header_footer(doc, d["source"], window)
    build_cover(doc, d)
    _page_break(doc)
    build_exec_summary(doc, d)
    _page_break(doc)
    build_volume_section(doc, d, charts_dir)
    build_action_mix(doc, d, charts_dir)
    _page_break(doc)
    build_user_section(doc, d, charts_dir)
    build_action_detail_section(doc, d, charts_dir)
    _page_break(doc)
    build_context_section(doc, d, charts_dir)
    build_recommendations(doc, d)
    _page_break(doc)
    build_methodology(doc, d)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"wrote {out_path}  ({out_path.stat().st_size:,} bytes)")
    return out_path


def main():
    ap = argparse.ArgumentParser(description=(
        "Build a CTO-facing DOCX report from a data.json artefact."))
    ap.add_argument("--data", required=True,
                    help="Path to reports/<slug>_<window>/data.json")
    args = ap.parse_args()
    build_report(Path(args.data))


if __name__ == "__main__":
    main()
